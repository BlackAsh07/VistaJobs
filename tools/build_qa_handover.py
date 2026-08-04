from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output" / "qa-handoff"
DOCX_PATH = OUT_DIR / "VistaJobs_QA_Handover_Document.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(28, 36, 52)
MUTED = RGBColor(90, 100, 116)
LIGHT_FILL = "F2F4F7"
BLUE_FILL = "E8EEF5"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run(run, size=11, bold=None, color=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run(run, size=16 if level == 1 else 13 if level == 2 else 12, bold=True, color=BLUE if level < 3 else DARK_BLUE)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.style = "Normal"
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run(r1, bold=True, color=INK)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run(r2, color=INK)
    else:
        set_run(p.add_run(text), color=INK)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        set_run(p.add_run(item), color=INK)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        set_run(p.add_run(item), color=INK)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_FILL):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_shading(hdr[i], header_fill)
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(header), bold=True, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if len(str(value)) > 18 else WD_ALIGN_PARAGRAPH.CENTER
            set_run(p.add_run(str(value)), color=INK)
    set_table_width(table, widths)
    return table


def add_metadata(doc, rows):
    for label, value in rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        set_run(p.add_run(f"{label}: "), bold=True, color=INK)
        set_run(p.add_run(value), color=INK)


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(header.add_run("VistaJobs QA Handover"), size=9, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(footer.add_run("Development-to-QA handover | Page "), size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run = footer.add_run()
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def build_doc():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("VISTAJOBS PLATFORM")
    set_run(run, size=24, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    set_run(subtitle.add_run("Software Development Document - Development to QA Handover"), size=14, color=MUTED)

    add_metadata(doc, [
        ("Project", "VistaJobs - recruitment, verification, employer matching and administration platform"),
        ("Prepared for", "QA Testing and TL Review"),
        ("Prepared by", "Development Team"),
        ("Version", "1.0"),
        ("Date", "July 2026"),
        ("Status", "Ready for QA validation after backend/frontend deployment"),
    ])

    add_heading(doc, "1. Purpose", 1)
    add_body(doc, "This document gives QA engineers a concise technical and functional overview of the VistaJobs platform so they can perform functional, integration, regression, security, responsive UI and UAT testing without reading the full codebase first.")

    add_heading(doc, "2. Application Overview", 1)
    add_body(doc, "VistaJobs is an ASP.NET Core Web API plus static frontend recruitment platform. It supports jobseeker registration, role-based login, fresher and experienced candidate profiles, Aadhaar/PAN/UAN verification, resume upload, employer registration with company validation, employer candidate matching, job postings, applications and admin dashboards.")

    add_heading(doc, "3. Technology Stack", 1)
    add_table(doc, ["Layer", "Technology / Implementation"], [
        ("Backend", "ASP.NET Core Web API (.NET 8), C#"),
        ("Frontend", "HTML, CSS, JavaScript static frontend under Frontend/"),
        ("Database", "SQL Server with Entity Framework Core migrations"),
        ("Authentication", "JWT bearer tokens, email OTP, BCrypt password hashing, optional Google sign-in"),
        ("Verification", "Local format validation plus DigiLocker/Surepass EPFO integration points"),
        ("Deployment", "Azure App Service API, Azure Static Apps or local Live Server frontend"),
    ], [2300, 7060])

    add_heading(doc, "4. Architecture", 1)
    add_body(doc, "Browser frontend -> API controllers -> services/providers -> Entity Framework Core -> SQL Server. Resume files are uploaded to the API and served through the /Uploads path. The frontend stores the authenticated session in sessionStorage and routes users by the role returned from login.")
    add_table(doc, ["Component", "Primary Files", "QA Focus"], [
        ("API startup", "Program.cs, appsettings*.json", "CORS, JWT configuration, uploads, Swagger and seed data"),
        ("Controllers", "Auth, Candidates, Jobs, Applications, Admin, Verification", "Validation, HTTP status codes, authorization and response shape"),
        ("Models/Data", "Models/*, ApplicationDbContext.cs, Migrations/*", "Schema mapping, duplicate prevention and update behavior"),
        ("Frontend", "Frontend/index.html, js/app.js, css/*", "Role routing, forms, toasts, responsive states and API calls"),
    ], [1700, 3300, 4360], BLUE_FILL)

    add_heading(doc, "5. User Roles", 1)
    add_table(doc, ["Role", "Expected Access", "Must Not Access"], [
        ("Jobseeker", "Register, login, create/update fresher or experienced profile, upload resume, verify Aadhaar/PAN/UAN, apply for jobs", "Employer matching dashboard, admin lists"),
        ("Employer", "Register through employer form, login, enter matching requirements, view matching candidates, create jobs", "Sensitive verification APIs outside allowed matching view, admin dashboards"),
        ("Admin", "Dashboard summary, raw users, jobs and applications lists", "No unauthorized access without admin role"),
    ], [1700, 4660, 3000])

    add_heading(doc, "6. Core Modules", 1)
    add_table(doc, ["Module", "Responsibility", "Key Test Areas"], [
        ("Authentication", "Register/login via email OTP, Google auth guard, JWT issuance", "Required fields, invalid credentials, OTP expiry, role routing"),
        ("Jobseeker Profiles", "Fresher/experienced candidate details and resume upload", "Create, update, validation, duplicate profile by email"),
        ("Verification", "Aadhaar, PAN, UAN and verification status", "Format validation, provider fallback, status persistence"),
        ("Employer Registration", "Company name, official email, GST, CIN, website and OTP", "Personal email blocking, domain/website match, OTP verification"),
        ("Employer Matching", "Requirement entry and skill-based candidate matching", "Skill aliases, candidate type, experience sorting, location preference"),
        ("Jobs and Applications", "Job posting and candidate application tracking", "CRUD, duplicate application blocking, confirmation email"),
        ("Administration", "User/job/application visibility for admins", "Role authorization and dashboard counts"),
    ], [1800, 3560, 4000], BLUE_FILL)

    doc.add_page_break()
    add_heading(doc, "7. End-to-End QA Workflows", 1)
    add_heading(doc, "7.1 Jobseeker Workflow", 2)
    add_numbered(doc, [
        "Register as a jobseeker using name, email, password and email OTP.",
        "Login with the same email and password, complete login OTP, and confirm jobseeker dashboard routing.",
        "Submit a fresher profile and confirm /api/Candidates/my-profile returns the saved profile.",
        "Update the profile as experienced and verify existing email updates the same candidate row instead of creating a duplicate.",
        "Upload PDF/DOC/DOCX resume and confirm the generated /Uploads path opens.",
        "Run Aadhaar, PAN and UAN verification flows and verify status indicators.",
        "Apply for a job and confirm duplicate application is blocked.",
    ])

    add_heading(doc, "7.2 Employer Workflow", 2)
    add_numbered(doc, [
        "Open Register as employer and verify only official company email domains are accepted.",
        "Complete company name, official email, GST, CIN, website, password and OTP.",
        "Login as employer and confirm employer matching dashboard opens by role.",
        "Enter job role, candidate type, required skills, minimum experience and preferred location.",
        "Confirm candidates with matching skills are shown and non-matching candidates are excluded.",
        "Create a job posting and verify admin dashboard/job APIs reflect it.",
    ])

    add_heading(doc, "7.3 Admin Workflow", 2)
    add_numbered(doc, [
        "Start the API and confirm the default admin account is created or updated.",
        "Login with the seeded admin account provided by the development team.",
        "Verify admin dashboard counts and raw lists for users, jobs and applications.",
        "Confirm non-admin roles cannot open admin APIs or admin frontend views.",
    ])

    add_heading(doc, "8. Critical Business Rules", 1)
    add_bullets(doc, [
        "Protected APIs require JWT authentication and role checks where applicable.",
        "Registration and login flows require OTP verification before account/session completion.",
        "Employer registration must reject personal email domains and validate company identity fields.",
        "Candidate profile is keyed by login email and should update, not duplicate, for the same user.",
        "Employer matching must use only relevant candidate profile fields and must not expose unrelated sensitive data.",
        "Duplicate job applications must be blocked by job ID and candidate email.",
        "Resume upload accepts only PDF, DOC and DOCX files.",
        "Configuration values such as JWT key, SMTP, database and provider keys must be present in app settings/user-secrets/Azure settings.",
    ])

    add_heading(doc, "9. QA Scope", 1)
    add_table(doc, ["Area", "In Scope Checks"], [
        ("Functional", "Registration, login, OTP, role routing, profiles, verification, matching, jobs, applications and admin dashboards"),
        ("Validation", "Required fields, email format, official email policy, GST/CIN length, website matching, password rules and upload types"),
        ("Integration", "SQL Server persistence, SMTP OTP/email, verification provider configuration, static frontend/API CORS"),
        ("Security", "JWT requirement, role isolation, no sensitive data in public pages, session clearing on logout"),
        ("Responsive UI", "Chrome, Edge and Firefox on desktop plus mobile width smoke tests"),
        ("Regression", "Retest login, profile save, employer matching and admin after every deployment"),
    ], [1800, 7560], LIGHT_FILL)

    add_heading(doc, "10. Environment Summary", 1)
    add_table(doc, ["Environment Item", "Value / QA Note"], [
        ("Local frontend", "http://127.0.0.1:5500/index.html"),
        ("Local API", "https://localhost:7250/api or http://localhost:5295/api based on launch profile"),
        ("Swagger", "https://localhost:7250/swagger"),
        ("Azure API", "Configured in frontend for deployed origin"),
        ("Database", "SQL Server via ConnectionStrings:DefaultConnection"),
        ("Secrets", "Use user-secrets locally and Azure App Service configuration in hosting. Do not share secrets in QA artifacts."),
        ("Migration check", "GET /__migrations"),
    ], [2400, 6960], BLUE_FILL)

    add_heading(doc, "11. Deployment and Configuration Checks", 1)
    add_bullets(doc, [
        "Confirm Jwt:Key, Jwt:Issuer and Jwt:Audience are configured before starting the API.",
        "Confirm CORS allows http://127.0.0.1:5500, http://localhost:5500 and the Azure Static Apps origin.",
        "Confirm the frontend API base URL points to local API for local testing and Azure API for deployed testing.",
        "Confirm Uploads/ is writable and uploaded resume links are reachable.",
        "Confirm SMTP credentials are valid before OTP, reset password and application email testing.",
        "Confirm Google OAuth authorized origins before enabling Google sign-in in deployed environments.",
    ])

    add_heading(doc, "12. High-Level Testing Checklist", 1)
    checklist = [
        "Role-based access and logout",
        "Jobseeker registration and login OTP",
        "Employer registration and official email validation",
        "Admin seed and admin dashboard",
        "Fresher and experienced profile save/update",
        "Resume upload and public link access",
        "Aadhaar, PAN and UAN verification",
        "Employer matching candidate results",
        "Job creation and application duplicate prevention",
        "Validation messages and toast behavior",
        "CORS and API error handling",
        "Responsive UI and browser compatibility",
        "Regression testing before release",
    ]
    add_table(doc, ["Done", "Checklist Item", "Notes"], [("[ ]", item, "") for item in checklist], [900, 4300, 4160])

    add_heading(doc, "13. Risks and Assumptions", 1)
    add_bullets(doc, [
        "Email OTP and notification testing depends on valid SMTP settings and mailbox access.",
        "Verification provider behavior depends on DigiLocker/Surepass settings; local fallback behavior must be documented for QA.",
        "Production-like testing should use isolated QA data and must not use real Aadhaar/PAN/UAN values unless approved.",
        "Azure deployment must update both API and frontend; deploying only the frontend can leave old CORS/API behavior active.",
        "Google sign-in requires the deployed origin to be added in Google Cloud OAuth settings.",
    ])

    add_heading(doc, "14. Version History and Approval", 1)
    add_table(doc, ["Version", "Date", "Description"], [
        ("1.0", "July 2026", "Initial VistaJobs Development-to-QA handover"),
        ("1.1", "Future", "Update after QA feedback, deployment changes or new module additions"),
    ], [1400, 1800, 6160], LIGHT_FILL)

    add_metadata(doc, [
        ("Prepared By", "Development Team"),
        ("Reviewed By", "____________________________"),
        ("Approved By", "____________________________"),
        ("QA Sign-off Date", "____________________________"),
    ])

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
